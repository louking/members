"""community_taxonomy - fetch Discourse taxonomy and build .docx report"""

# standard
import json
from datetime import datetime

# pypi
from fluent_discourse import Discourse, DiscourseError
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# homegrown
from .community import (
    _RateLimitedDiscourse, _RateLimiter,
    DISCOURSE_RATE_LIMIT_MAX_CALLS, DISCOURSE_RATE_LIMIT_WINDOW_SECS,
)


# ---------------------------------------------------------------------------
# Data fetchers
# ---------------------------------------------------------------------------

def fetch_categories(discourse):
    data = discourse.categories.json.get({'include_subcategories': 'true'})
    cats = data.get('category_list', {}).get('categories', [])
    flat = []
    for c in cats:
        flat.append(c)
        for sub in c.get('subcategory_list', []):
            sub['_parent_name'] = c['name']
            flat.append(sub)
    return flat


def fetch_category_groups(discourse, query_id):
    """Run a Data Explorer query that returns category_id, group_name, permission_type rows."""
    resp = discourse.admin.plugins.explorer.queries._(query_id).run.post({'params': {}})
    columns = resp.get('columns', [])
    rows = resp.get('rows', [])
    return [dict(zip(columns, row)) for row in rows]


def fetch_tags(discourse):
    return discourse.tags.json.get({}).get('tags', [])


def fetch_tag_groups(discourse):
    return discourse.tag_groups.json.get({}).get('tag_groups', [])


def fetch_groups(discourse):
    groups = []
    page = 0
    while True:
        batch = discourse.groups.json.get({'page': page}).get('groups', [])
        if not batch:
            break
        groups.extend(batch)
        page += 1
        if len(batch) < 20:
            break
    return groups


def fetch_site_settings(discourse):
    data = discourse.admin.site_settings.json.get({})
    return {
        item['setting']: {
            'value': item.get('value', ''),
            'label': item.get('humanized_name', item['setting']),
        }
        for item in data.get('site_settings', [])
    }


def fetch_user_fields(discourse):
    return discourse.admin.user_fields.json.get({}).get('user_fields', [])


def fetch_badges(discourse):
    return discourse.admin.badges.json.get({}).get('badges', [])


def fetch_themes(discourse):
    return discourse.admin.themes.json.get({}).get('themes', [])


def fetch_watched_words(discourse):
    return discourse.admin.watched_words.json.get({})


def fetch_nav_items(discourse):
    try:
        return discourse.navigation_menu_items.json.get({}).get('navigation_menu_items', [])
    except (DiscourseError, Exception):
        return []


def fetch_site_title(discourse):
    try:
        return discourse.about.json.get({}).get('about', {}).get('title', '')
    except Exception:
        return ''


def fetch_all(base_url, api_key, api_username, category_groups_query_id=None):
    """Fetch all taxonomy data and return as a dict. Tolerates per-section errors."""
    discourse = _RateLimitedDiscourse(
        Discourse(
            base_url=base_url,
            username=api_username,
            api_key=api_key,
            raise_for_rate_limit=False,
        ),
        _RateLimiter(max_calls=DISCOURSE_RATE_LIMIT_MAX_CALLS, window_secs=DISCOURSE_RATE_LIMIT_WINDOW_SECS),
    )

    steps = [
        ('Groups',        'groups',        lambda: fetch_groups(discourse)),
        ('Categories',    'categories',    lambda: fetch_categories(discourse)),
        ('Tags',          'tags',          lambda: fetch_tags(discourse)),
        ('Tag groups',    'tag_groups',    lambda: fetch_tag_groups(discourse)),
        ('Site settings', 'site_settings', lambda: fetch_site_settings(discourse)),
        ('User fields',   'user_fields',   lambda: fetch_user_fields(discourse)),
        ('Badges',        'badges',        lambda: fetch_badges(discourse)),
        ('Themes',        'themes',        lambda: fetch_themes(discourse)),
        ('Watched words', 'watched_words', lambda: fetch_watched_words(discourse)),
        ('Nav items',     'nav_items',     lambda: fetch_nav_items(discourse)),
    ]

    data = {}
    for label, key, fn in steps:
        print(f'  Fetching {label}...', end=' ', flush=True)
        try:
            data[key] = fn()
            count = len(data[key]) if isinstance(data[key], list) else len(data[key])
            print(f'OK ({count} items)')
        except Exception as e:
            print(f'WARN: {e}')
            data[key] = [] if key != 'site_settings' else {}

    data['title'] = fetch_site_title(discourse)

    # Merge category group permissions from Data Explorer query
    if category_groups_query_id and data.get('categories'):
        print('  Fetching category group permissions...', end=' ', flush=True)
        try:
            cg_rows = fetch_category_groups(discourse, category_groups_query_id)
            perms_by_cat = {}
            for row in cg_rows:
                perms_by_cat.setdefault(row['category_id'], []).append({
                    'group_name': row['group_name'],
                    'permission_type': row['permission_type'],
                })
            for c in data['categories']:
                if c['id'] in perms_by_cat:
                    c['group_permissions'] = perms_by_cat[c['id']]
            print(f'OK ({len(cg_rows)} rows)')
        except Exception as e:
            print(f'WARN: {e}')

    return data


# ---------------------------------------------------------------------------
# Site settings filter list
# ---------------------------------------------------------------------------

SETTINGS_OF_INTEREST = [
    'min_trust_level_to_create_tag',
    'min_trust_to_create_tag',
    'min_trust_level_to_tag_topics',
    'min_trust_to_post_links',
    'min_trust_level_to_allow_invite',
    'email_in_min_trust',
    'login_required',
    'enable_google_oauth2_logins',
    'enable_local_logins',
    'sso_url',
    'discourse_connect_url',
    'enable_discourse_connect',
    'must_approve_users',
    'invite_only',
    'allow_new_registrations',
    'enable_personal_messages',
    'min_trust_level_to_send_messages',
    'personal_message_enabled_groups',
    'rate_limit_create_topic',
    'rate_limit_new_user_create_topic',
    'max_topics_per_day',
    'tagging_enabled',
    'max_tags_per_topic',
    'max_tag_length',
    'tag_topic_allowed_groups',
    'fixed_category_positions',
    'allow_uncategorized_topics',
    'notification_email',
    'reply_by_email_enabled',
    'email_prefix',
    'title',
    'site_description',
    'short_site_description',
    'contact_email',
    'contact_url',
]


# ---------------------------------------------------------------------------
# Document styling helpers
# ---------------------------------------------------------------------------

HEADER_COLOR = RGBColor(0x2E, 0x55, 0x9A)


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = HEADER_COLOR
    return p


def _add_table(doc, headers, rows, col_widths_inches=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(cell, '2E559A')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = 'F0F4FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val) if val is not None else ''
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            _set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths_inches:
        for i, width in enumerate(col_widths_inches):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


def _permission_label(ptype):
    return {1: 'Full', 2: 'Create Post', 3: 'Read Only'}.get(ptype, str(ptype))


def _visibility_label(v):
    return {0: 'Public', 1: 'Logged-in users', 2: 'Members', 3: 'Staff', 4: 'Owners'}.get(v, str(v))


def _access_level_label(v):
    return {
        0: 'Nobody',
        1: 'Admins only',
        2: 'Admins & Mods',
        3: 'Members, Admins & Mods',
        4: 'Owners, Admins & Mods',
        99: 'Everyone',
    }.get(v, str(v))


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def _section_cover(doc, base_url, title='Discourse Community Forum'):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title or 'Discourse Community Forum')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = HEADER_COLOR

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Taxonomy & Configuration Reference')
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f'Source: {base_url}')
    r4.font.size = Pt(10)
    r4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def _section_groups(doc, groups):
    _add_heading(doc, 'Groups', level=1)
    doc.add_paragraph(
        'Discourse groups control category access, permissions, and trust level grants. '
        'System groups (trust_level_0 through trust_level_4, admins, moderators, staff) '
        'are managed automatically.'
    )

    system_names = {'admins', 'moderators', 'staff', 'trust_level_0', 'trust_level_1',
                    'trust_level_2', 'trust_level_3', 'trust_level_4'}
    custom = [g for g in groups if g.get('name') not in system_names]
    system = [g for g in groups if g.get('name') in system_names]

    _add_heading(doc, 'Custom Groups', level=2)
    if custom:
        rows = []
        for g in sorted(custom, key=lambda x: x.get('name', '')):
            rows.append([
                g.get('name', ''),
                g.get('full_name', ''),
                _visibility_label(g.get('visibility_level', '')),
                g.get('user_count', ''),
                _access_level_label(g.get('mentionable_level', '')),
                _access_level_label(g.get('messageable_level', '')),
            ])
        _add_table(doc,
                   ['Name', 'Full Name', 'Visibility', 'Members', 'Mentionable', 'Messageable'],
                   rows,
                   col_widths_inches=[1.6, 1.8, 1.2, 0.7, 1.1, 1.1])
    else:
        doc.add_paragraph('No custom groups found.')

    _add_heading(doc, 'System Groups', level=2)
    rows = [[g.get('name', ''), g.get('user_count', '')] for g in system]
    _add_table(doc, ['Name', 'Members'], rows, col_widths_inches=[3.0, 1.0])
    doc.add_paragraph()


def _section_categories(doc, categories):
    _add_heading(doc, 'Categories', level=1)
    doc.add_paragraph(
        'Categories organize forum content. Subcategories are indented under their parent. '
        'Permissions show which groups have access and at what level (Full / Create Post / Read Only).'
    )

    BASE_INDENT = Inches(0.25)
    SUB_INDENT = Inches(0.6)
    DETAIL_EXTRA = Inches(0.2)

    # Sort: public before restricted, alphabetical within each; subcategories follow their parent
    def _cat_sort_key(c):
        return (c.get('read_restricted', False), c.get('name', '').lower())

    top_level = sorted([c for c in categories if not c.get('_parent_name')], key=_cat_sort_key)
    subs_by_parent = {}
    for c in categories:
        if c.get('_parent_name'):
            subs_by_parent.setdefault(c['_parent_name'], []).append(c)
    for lst in subs_by_parent.values():
        lst.sort(key=_cat_sort_key)

    ordered = []
    for c in top_level:
        ordered.append(c)
        ordered.extend(subs_by_parent.get(c.get('name', ''), []))

    for c in ordered:
        is_sub = bool(c.get('_parent_name', ''))
        indent = SUB_INDENT if is_sub else BASE_INDENT
        name = c.get('name', '')
        description = (c.get('description_text', '') or '').strip()

        group_perms = c.get('group_permissions', [])
        if group_perms:
            perm_str = ', '.join(
                f'{gp["group_name"]}: {_permission_label(gp["permission_type"])}'
                for gp in group_perms
            )
        elif c.get('read_restricted'):
            perm_str = 'Restricted'
        else:
            perm_str = 'Public'

        # Name line (no slug)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = indent
        p.paragraph_format.space_after = Pt(0)
        name_run = p.add_run(('⮣ ' if is_sub else '• ') + name)
        name_run.bold = True
        name_run.font.size = Pt(10)

        # Description line (before access)
        if description:
            pd = doc.add_paragraph()
            pd.paragraph_format.left_indent = indent + DETAIL_EXTRA
            pd.paragraph_format.space_after = Pt(0)
            desc_run = pd.add_run(description)
            desc_run.font.size = Pt(9)
            desc_run.font.italic = True
            desc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Access — label then one line per group (or single line if public/restricted)
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = indent + DETAIL_EXTRA
        pa.paragraph_format.space_after = Pt(0)
        access_label = pa.add_run('Access: ')
        access_label.font.size = Pt(9)
        access_label.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        access_label.bold = True

        if group_perms:
            pa.paragraph_format.space_after = Pt(0)
            for gp in group_perms:
                pg = doc.add_paragraph()
                pg.paragraph_format.left_indent = indent + DETAIL_EXTRA + Inches(0.2)
                pg.paragraph_format.space_after = Pt(0)
                gr = pg.add_run(f'• {gp["group_name"]}: {_permission_label(gp["permission_type"])}')
                gr.font.size = Pt(9)
                gr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            doc.paragraphs[-1].paragraph_format.space_after = Pt(5)
        else:
            access_label.bold = False
            access_label_rest = pa.add_run(perm_str)
            access_label_rest.font.size = Pt(9)
            access_label_rest.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            pa.paragraph_format.space_after = Pt(5)

    doc.add_paragraph()


def _section_tags(doc, tags, tag_groups, group_id_to_name=None):
    _add_heading(doc, 'Tags', level=1)

    _add_heading(doc, 'Tag Groups', level=2)
    doc.add_paragraph('Tag groups organize tags and can restrict which groups may use them.')
    if tag_groups:
        INDENT = Inches(0.25)
        DETAIL = Inches(0.45)
        for tg in tag_groups:
            tag_names = ', '.join(t['name'] for t in tg.get('tags', []))
            one_per = 'Yes' if tg.get('one_per_topic') else 'No'
            perms = tg.get('permissions') or {}
            if perms and group_id_to_name:
                perm_str = ', '.join(
                    f'{group_id_to_name.get(int(gid), gid)}: {_permission_label(ptype)}'
                    for gid, ptype in perms.items()
                )
            elif perms:
                perm_str = ', '.join(f'group {gid}: {_permission_label(ptype)}' for gid, ptype in perms.items())
            else:
                perm_str = 'All'

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = INDENT
            p.paragraph_format.space_after = Pt(0)
            name_run = p.add_run(f'• {tg.get("name", "")}')
            name_run.bold = True
            name_run.font.size = Pt(10)

            for label, value in [('Tags', tag_names or 'none'), ('Access', perm_str), ('One per topic', one_per)]:
                pd = doc.add_paragraph()
                pd.paragraph_format.left_indent = DETAIL
                pd.paragraph_format.space_after = Pt(0)
                lr = pd.add_run(f'{label}: ')
                lr.font.size = Pt(9)
                lr.bold = True
                vr = pd.add_run(value)
                vr.font.size = Pt(9)

            doc.paragraphs[-1].paragraph_format.space_after = Pt(5)
    else:
        doc.add_paragraph('No tag groups configured.')

    _add_heading(doc, 'All Tags', level=2)
    if tags:
        for name in sorted(t.get('name', '') for t in tags):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(0)
            p.add_run(f'• {name}').font.size = Pt(9)
    else:
        doc.add_paragraph('No tags found.')
    doc.add_paragraph()


def _section_site_settings(doc, settings):
    _add_heading(doc, 'Site Settings (Selected)', level=1)
    doc.add_paragraph(
        'Key policy and configuration settings. Full settings are available in Admin > Settings.'
    )

    rows = []
    for key in SETTINGS_OF_INTEREST:
        if key in settings:
            entry = settings[key]
            rows.append([entry['label'], entry['value']])

    for key, entry in settings.items():
        if ('connect' in key or 'sso' in key) and key not in SETTINGS_OF_INTEREST:
            rows.append([entry['label'], entry['value']])

    if rows:
        _add_table(doc, ['Setting', 'Value'], rows, col_widths_inches=[3.5, 5.8])
    else:
        doc.add_paragraph('Could not retrieve site settings.')
    doc.add_paragraph()


def _section_user_fields(doc, user_fields):
    _add_heading(doc, 'Custom User Fields', level=1)
    doc.add_paragraph(
        'Custom profile fields added to Discourse, beyond built-in username/email/name.'
    )
    if user_fields:
        rows = []
        for f in user_fields:
            rows.append([
                str(f.get('id', '')),
                f.get('name', ''),
                f.get('field_type', ''),
                'Yes' if f.get('required') else 'No',
                'Yes' if f.get('show_on_profile') else 'No',
                'Yes' if f.get('show_on_user_card') else 'No',
                'Yes' if f.get('editable') else 'No',
            ])
        _add_table(doc,
                   ['ID', 'Name', 'Type', 'Required', 'On Profile', 'On Card', 'Editable'],
                   rows,
                   col_widths_inches=[0.4, 2.0, 1.0, 0.7, 0.8, 0.7, 0.7])
    else:
        doc.add_paragraph('No custom user fields configured.')
    doc.add_paragraph()


def _section_badges(doc, badges):
    _add_heading(doc, 'Badges', level=1)
    doc.add_paragraph('Active badges in use on the forum.')
    enabled = [b for b in badges if b.get('enabled')]
    if enabled:
        rows = []
        for b in sorted(enabled, key=lambda x: x.get('name', '')):
            rows.append([
                b.get('name', ''),
                b.get('badge_type', {}).get('name', '') if isinstance(b.get('badge_type'), dict) else '',
                'Yes' if b.get('system') else 'No',
                'Yes' if b.get('manually_grantable') else 'No',
                str(b.get('grant_count', 0)),
            ])
        _add_table(doc, ['Name', 'Type', 'System', 'Manual', 'Granted'],
                   rows, col_widths_inches=[2.5, 1.2, 0.7, 0.7, 0.8])
    else:
        doc.add_paragraph('No enabled badges found.')
    doc.add_paragraph()


def _section_themes(doc, themes):
    _add_heading(doc, 'Themes & Components', level=1)
    doc.add_paragraph(
        'Installed themes and components. The active theme controls the forum\'s appearance.'
    )
    if themes:
        rows = []
        for t in sorted(themes, key=lambda x: (not x.get('default'), x.get('name', ''))):
            rows.append([
                t.get('name', ''),
                'Yes' if t.get('default') else '',
                'Component' if t.get('component') else 'Theme',
                'Yes' if t.get('enabled') else 'No',
                str(t.get('user_selectable', '')),
            ])
        _add_table(doc, ['Name', 'Active', 'Type', 'Enabled', 'User Selectable'],
                   rows, col_widths_inches=[3.0, 0.7, 1.0, 0.8, 1.2])
    else:
        doc.add_paragraph('No themes found.')
    doc.add_paragraph()


def _section_watched_words(doc, words):
    _add_heading(doc, 'Watched Words', level=1)
    doc.add_paragraph(
        'Words configured for automatic action (block, require approval, flag, replace, link).'
    )
    if words:
        by_action = {}
        for w in words:
            action = w.get('action', 'unknown')
            by_action.setdefault(action, []).append(w.get('word', ''))
        for action, word_list in sorted(by_action.items()):
            _add_heading(doc, f'Action: {action}', level=2)
            rows = [[w] for w in sorted(word_list)]
            _add_table(doc, ['Word / Pattern'], rows, col_widths_inches=[9.0])
    else:
        doc.add_paragraph('No watched words configured.')
    doc.add_paragraph()


def _section_nav(doc, items):
    _add_heading(doc, 'Navigation Menu Items', level=1)
    doc.add_paragraph('Custom items added to the forum\'s top navigation menu.')
    if items:
        rows = []
        for item in items:
            rows.append([
                item.get('name', ''),
                item.get('url', ''),
                'Yes' if item.get('target_blank') else 'No',
            ])
        _add_table(doc, ['Label', 'URL', 'New Tab'], rows, col_widths_inches=[2.0, 6.0, 0.8])
    else:
        doc.add_paragraph('No custom navigation items configured (or endpoint not available).')
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def build_docx(data, base_url):
    """Build and return a python-docx Document from pre-fetched taxonomy data."""
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    _section_cover(doc, base_url, title=data.get('title', ''))
    _section_groups(doc, data['groups'])
    doc.add_page_break()
    _section_categories(doc, data['categories'])
    doc.add_page_break()
    group_id_to_name = {g['id']: g['name'] for g in data.get('groups', [])}
    _section_tags(doc, data['tags'], data['tag_groups'], group_id_to_name=group_id_to_name)
    doc.add_page_break()
    _section_site_settings(doc, data['site_settings'])
    doc.add_page_break()
    _section_user_fields(doc, data['user_fields'])
    _section_badges(doc, data['badges'])
    doc.add_page_break()
    _section_themes(doc, data['themes'])
    _section_watched_words(doc, data['watched_words'])
    _section_nav(doc, data['nav_items'])

    return doc
